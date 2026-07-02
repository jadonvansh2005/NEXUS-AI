from typing import List

from docx import Document

from rag.loaders.base_loader import (
    BaseLoader
)

from rag.loaders.loader_constants import (
    DOCX_EXTENSIONS,
    DocumentType
)

from rag.loaders.loader_factory import (
    LoaderFactory
)

from rag.loaders.loader_schema import (
    LoadedDocument
)

from rag.loaders.loader_utils import (
    LoaderUtils
)


class DOCXLoader(

    BaseLoader

):

    SUPPORTED_EXTENSIONS = (

        DOCX_EXTENSIONS

    )

    def load(

        self

    ) -> LoadedDocument:

        document = (

            Document(

                self.absolute_path

            )

        )

        paragraphs: List[str] = []

        for paragraph in (

            document.paragraphs

        ):

            text = (

                paragraph.text

            ).strip()

            if text:

                paragraphs.append(

                    LoaderUtils.normalize_text(

                        text

                    )

                )

        full_text = (

            "\n\n".join(

                paragraphs

            )

        )

        return LoadedDocument(

            text=full_text,

            document_type=(

                DocumentType.DOCX.value

            ),

            source_path=(

                self.absolute_path

            ),

            filename=(

                self.filename

            ),

            metadata=self.metadata(),

            page_count=1,

            pages=[

                full_text

            ],

            checksum=(

                LoaderUtils.calculate_checksum(

                    self.absolute_path

                )

            ),

            file_size=(

                self.filesize

            )

        )


LoaderFactory.register(

    DOCXLoader

)